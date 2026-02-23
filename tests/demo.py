"""
Document Ingestion Pipeline — Live Demo
Run: python demo.py
"""

import sys, textwrap
sys.path.insert(0, "/home/claude")

from doc_extracter import DocumentIngestionPipeline, PageType


SEP  = "═" * 68
SEP2 = "─" * 68


def banner(title: str):
    print(f"\n{SEP}")
    print(f"  {title}")
    print(SEP)


def show_result(result, max_chars: int = 500):
    print(result.summary())
    print(f"\n{SEP2}")
    print("  Page breakdown:")
    print(SEP2)
    for page in result.pages:
        tag   = f"[{page.page_type.value.upper():12s}]"
        conf  = f"  conf={page.confidence:.0f}%" if page.confidence else ""
        words = f"  {page.word_count} words"
        label = f"  Page {page.page_number}{words}{conf}"
        snippet = page.text[:80].replace("\n", " ")
        print(f"  {tag}{label}")
        print(f"    ↳ {snippet!r}")
    print(f"\n{SEP2}")
    print("  Full text preview:")
    print(SEP2)
    text = result.plain_text
    print(textwrap.fill(text[:max_chars], width=65, subsequent_indent="  "))
    if len(text) > max_chars:
        print(f"  … ({len(text) - max_chars} more chars)")


# ── Create test files ──────────────────────────────────────────────────────

def make_test_files():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from PIL import Image, ImageDraw
    from docx import Document

    # 1. Native-text PDF
    c = canvas.Canvas("data/demo/demo_native.pdf", pagesize=letter)
    w, h = letter
    c.setFont("Helvetica-Bold", 15); c.drawString(72, h-80, "SERVICE AGREEMENT")
    c.setFont("Helvetica", 11)
    c.drawString(72, h-120, "1. SCOPE OF SERVICES")
    c.drawString(72, h-145, "Provider shall deliver software consulting services as defined herein.")
    c.drawString(72, h-165, "Services include design, implementation, and post-launch support.")
    c.showPage()
    c.setFont("Helvetica-Bold", 13); c.drawString(72, h-80, "2. PAYMENT TERMS")
    c.setFont("Helvetica", 11)
    c.drawString(72, h-110, "Invoices are due within 30 days of issuance. Late fees apply.")
    c.drawString(72, h-130, "Preferred payment method: bank transfer or corporate cheque.")
    c.save()

    # 2. Scanned PDF (image pages)
    def make_page_img(lines):
        img = Image.new("RGB", (780, 480), "white")
        draw = ImageDraw.Draw(img)
        y = 35
        for line in lines:
            draw.text((40, y), line, fill="black")
            y += 30
        return img

    img1 = make_page_img([
        "LEASE AGREEMENT — PAGE 1",
        "",
        "Article 1 — Parties",
        "(1) Landlord: Evergreen Realty LLC",
        "(2) Tenant:   John Smith",
        "",
        "Article 2 — Premises",
        "The property at 45 Oak Avenue, Boston MA 02101",
        "is leased for residential use only.",
    ])
    img2 = make_page_img([
        "LEASE AGREEMENT — PAGE 2",
        "",
        "Article 3 — Rent",
        "(1) Monthly rent: USD 2,200 due on the 1st.",
        "(2) Grace period: 5 calendar days.",
        "",
        "Article 4 — Termination",
        "(1) Either party may terminate with 60-day notice.",
    ])
    img1.save("data/demo/p1.png"); img2.save("data/demo/p2.png")
    c2 = canvas.Canvas("data/demo/demo_scanned.pdf", pagesize=letter)
    c2.drawImage("data/demo/p1.png", 0, 0, width=w, height=h)
    c2.showPage()
    c2.drawImage("data/demo/p2.png", 0, 0, width=w, height=h)
    c2.save()

    # 3. DOCX
    doc = Document()
    doc.add_heading("PRIVACY POLICY", level=1)
    doc.add_heading("1. Data We Collect", level=2)
    doc.add_paragraph("We collect personal data necessary for service delivery:")
    doc.add_paragraph("Name, email, and payment information.", style="List Bullet")
    doc.add_paragraph("Device identifiers and usage logs.",  style="List Bullet")
    doc.add_heading("2. Your Rights", level=2)
    doc.add_paragraph(
        "You may request access, correction, or deletion of your data "
        "at any time by contacting our Data Protection Officer."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Right"
    table.rows[0].cells[1].text = "How to Exercise"
    for right, method in [
        ("Access", "Email dpo@company.com"),
        ("Deletion", "Use the in-app settings"),
        ("Portability", "Download via account page"),
    ]:
        r = table.add_row().cells
        r[0].text = right; r[1].text = method
    doc.save("data/demo/demo_doc.docx")

    # 4. PNG image
    img = Image.new("RGB", (860, 320), "white")
    draw = ImageDraw.Draw(img)
    for i, line in enumerate([
        "TERMS OF SERVICE — SUMMARY",
        "",
        "By using this service, you agree to the following:",
        "  1. You will not use this service for illegal purposes.",
        "  2. You are responsible for maintaining your account security.",
        "  3. We reserve the right to update these terms with 30-day notice.",
    ]):
        draw.text((40, 30 + i * 38), line, fill="black")
    img.save("data/demo/demo_image.png")

    # 5. Text file
    with open("data/demo/demo_text.txt", "w") as f:
        f.write("""MEETING MINUTES — BOARD MEETING
Date: August 15, 2025

ATTENDEES
Alice Chen (Chair), Bob Martinez (CFO), Carol Singh (CTO)

AGENDA ITEM 1 — Q3 Financial Review
Q3 revenue came in at $4.2M, 12% above forecast.
Operating expenses were well controlled at $2.8M.

AGENDA ITEM 2 — Product Roadmap
The team approved the v5.0 feature set for Q4 release.
Key features: AI assistant integration, multi-language support.

ACTION ITEMS
1. Alice: Publish Q3 earnings report by August 20.
2. Bob:   Finalise Q4 budget by September 1.
3. Carol: Share v5.0 spec for board review.
""")

    print("All test files created.\n")


# ── Run demo ──────────────────────────────────────────────────────────────

def main():
    make_test_files()

    pipeline = DocumentIngestionPipeline(ocr_dpi=300)

    print(f"\n  Supported formats: {pipeline.supported_formats}")

    test_cases = [
        ("data/demo/demo_native.pdf",  "1.  Native Text PDF          (2 pages, no OCR needed)"),
        ("data/demo/demo_scanned.pdf", "2.  Scanned PDF              (2 image pages → OCR)"),
        ("data/demo/demo_doc.docx",    "3.  Word Document .docx      (headings + table)"),
        ("data/demo/demo_image.png",   "4.  PNG Image                (standalone image → OCR)"),
        ("data/demo/demo_text.txt",    "5.  Plain Text .txt          (meeting minutes)"),
        ("data/demo/nonexistent.pdf",  "6.  Missing File             (error handling)"),
    ]

    for path, label in test_cases:
        banner(label)
        result = pipeline.ingest(path)
        show_result(result)

    # ── Batch mode demo ───────────────────────────────────────────────────
    banner("7.  Batch Processing (all 5 real files)")
    results = pipeline.ingest_batch([
        "data/demo/demo_native.pdf",
        "data/demo/demo_scanned.pdf",
        "data/demo/demo_doc.docx",
        "data/demo/demo_image.png",
        "data/demo/demo_text.txt",
    ])
    print(f"\n  Processed {len(results)} documents\n")
    for r in results:
        from pathlib import Path
        name  = Path(r.file_path).name
        pages = r.total_pages
        ocr   = r.ocr_page_count
        nat   = r.native_page_count
        conf  = f"  OCR conf={r.avg_ocr_confidence:.0f}%" if r.avg_ocr_confidence else ""
        status = "✓" if r.status.value == "success" else "!"
        print(f"  [{status}] {name:<28} {pages}p  native={nat}  ocr={ocr}{conf}")

    print(f"\n{SEP}\n  Demo complete.\n{SEP}\n")


if __name__ == "__main__":
    main()